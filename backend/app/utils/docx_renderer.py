"""DOCX to HTML conversion utility."""

import html
import logging
from pathlib import Path

from fastapi import HTTPException

logger = logging.getLogger(__name__)


def convert_docx_to_html(file_path: Path) -> str:
    """Convert a .docx file to HTML content.

    Extracts paragraphs and tables from the DOCX and renders them
    as semantic HTML elements.

    Args:
        file_path: Path to the .docx file.

    Returns:
        HTML string representing the document content.

    Raises:
        HTTPException(422): If the DOCX file is corrupted or uses unsupported features.
    """
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx library not available",
        )

    try:
        doc = Document(str(file_path))
    except (PackageNotFoundError, Exception) as e:
        logger.error("Failed to open DOCX file %s: %s", file_path, e)
        raise HTTPException(
            status_code=422,
            detail=f"Unable to process DOCX file: {e}",
        )

    parts: list[str] = []

    try:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name.lower() if paragraph.style else ""

            if "heading 1" in style_name:
                parts.append(f"<h1>{html.escape(text)}</h1>")
            elif "heading 2" in style_name:
                parts.append(f"<h2>{html.escape(text)}</h2>")
            elif "heading 3" in style_name:
                parts.append(f"<h3>{html.escape(text)}</h3>")
            elif "heading" in style_name:
                parts.append(f"<h4>{html.escape(text)}</h4>")
            else:
                parts.append(f"<p>{html.escape(text)}</p>")

        # Process tables
        for table in doc.tables:
            parts.append("<table border='1' cellpadding='4' cellspacing='0'>")
            for row_idx, row in enumerate(table.rows):
                parts.append("<tr>")
                for cell in row.cells:
                    tag = "th" if row_idx == 0 else "td"
                    parts.append(f"<{tag}>{html.escape(cell.text)}</{tag}>")
                parts.append("</tr>")
            parts.append("</table>")

    except Exception as e:
        logger.error("Error processing DOCX content from %s: %s", file_path, e)
        raise HTTPException(
            status_code=422,
            detail=f"Error processing DOCX content: {e}",
        )

    return "\n".join(parts)
